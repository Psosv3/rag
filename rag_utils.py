import os
import faiss
import numpy as np
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from PyPDF2 import PdfReader
import docx
from dotenv import load_dotenv
from typing import Optional, Union, List, Dict
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from langchain.schema import BaseRetriever, Document
from langchain.base_language import BaseLanguageModel
import models
import uuid
from pathlib import Path
from tqdm.auto import tqdm
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.docstore.in_memory import InMemoryDocstore

from flashrank import Ranker
from langchain_community.document_compressors import FlashrankRerank
from langchain.retrievers.contextual_compression import ContextualCompressionRetriever

#---------------------------------
from mistralai import Mistral

client = Mistral(api_key = models.mistral_api_key)
#---------------------------------

# Load environment variables
load_dotenv()

# Stockage global des vectorstores par entreprise
vectorstores_cache = {}

def get_company_data_dir(company_id: str, base_data_dir: str = "data") -> str:
    """Retourne le répertoire de données pour une entreprise spécifique."""
    company_dir = os.path.join(base_data_dir, f"company_{company_id}")
    os.makedirs(company_dir, exist_ok=True)
    return company_dir

def get_company_index_dir(company_id: str, base_data_dir: str = "data") -> str:
    """Retourne le répertoire d'index pour une entreprise spécifique."""
    index_dir = os.path.join(base_data_dir, f"indexes", f"company_{company_id}")
    os.makedirs(index_dir, exist_ok=True)
    return index_dir

def read_pdf(file_path):
    """Reads a PDF file and returns its text content."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Erreur lors de la lecture du PDF {file_path}: {str(e)}")
        return ""

def read_docx(file_path):
    """Reads a DOCX file and returns its text content."""
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Erreur lors de la lecture du DOCX {file_path}: {str(e)}")
        return ""

def load_documents(data_dir):
    """Loads all PDF and DOCX documents from a directory."""
    docs = []
    for fname in os.listdir(data_dir):
        fpath = os.path.join(data_dir, fname)
        if fname.endswith(".pdf"):
            docs.append(read_pdf(fpath))
        elif fname.endswith(".docx"):
            docs.append(read_docx(fpath))
    return docs

def create_vectorstore(docs: List[str],
                      *,
                      model: str = "text-embedding-3-large",   # Use OpenAI, or swap for multilingual if needed
                      splitter_chunk_size: int = 800,
                      splitter_overlap: int = 100,
                      embed_batch_size: int = 128,
                      use_hnsw: bool = True,
                      hnsw_m: int = 32,
                      normalise: bool = True,
                      persist_dir: Optional[str | Path] = None,
                      show_progress: bool = True,
                      ) -> FAISS:
    """
    Create and optionally persist a FAISS vector index from document strings.
    """
    if not docs or all(not d.strip() for d in docs):
        raise ValueError("No non-empty documents provided.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=splitter_chunk_size,
        chunk_overlap=splitter_overlap,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    documents: List[Document] = []
    for src_id, raw in enumerate(docs):
        if not raw.strip():
            continue
        for chunk_id, chunk in enumerate(splitter.split_text(raw)):
            meta = {"source_id": src_id, "chunk_id": chunk_id}
            documents.append(Document(page_content=chunk, metadata=meta))

    if not documents:
        raise ValueError("Everything was filtered out during chunking.")

    embedder = OpenAIEmbeddings(
        model=model,
        chunk_size=embed_batch_size,
        show_progress_bar=show_progress,
        max_retries=6,
    )

    texts = [d.page_content for d in documents]
    vectors = embedder.embed_documents(texts)

    vecs_np = np.asarray(vectors, dtype="float32")
    if normalise:
        faiss.normalize_L2(vecs_np)

    dimension = vecs_np.shape[1]

    if use_hnsw:
        index = faiss.IndexHNSWFlat(dimension, hnsw_m)
        index.hnsw.efConstruction = max(64, hnsw_m * 4)
        index.hnsw.efSearch = 128
    else:
        index = faiss.IndexFlatIP(dimension)

    index.add(vecs_np)

    ids = [str(uuid.uuid4()) for _ in documents]
    docstore = InMemoryDocstore(dict(zip(ids, documents)))
    index_to_docstore_id = {i: doc_id for i, doc_id in enumerate(ids)}

    vectordb = FAISS(
        embedding_function=embedder,
        index=index,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
    )

    if persist_dir:
        Path(persist_dir).mkdir(parents=True, exist_ok=True)
        vectordb.save_local(str(persist_dir))

    return vectordb

def build_index(company_id: str, data_dir: str = "data", HTTPException=None):
    """Builds the vectorstore index from documents for a specific company."""
    try:
        # Répertoire spécifique à l'entreprise
        company_data_dir = get_company_data_dir(company_id, data_dir)
        company_index_dir = get_company_index_dir(company_id, data_dir)
        
        # Charger les documents de l'entreprise
        docs = load_documents(company_data_dir)
        if not docs:
            if HTTPException:
                raise HTTPException(status_code=400, detail=f"Aucun document trouvé pour l'entreprise {company_id}")
            else:
                raise ValueError(f"Aucun document trouvé pour l'entreprise {company_id}")
        
        # Créer le vectorstore avec persistance
        vectordb = create_vectorstore(docs, persist_dir=company_index_dir)
        
        # Mettre en cache
        vectorstores_cache[company_id] = vectordb
        
        return vectordb
    except Exception as e:
        if HTTPException:
            raise HTTPException(status_code=500, detail=f"Erreur lors de la construction de l'index: {str(e)}")
        else:
            raise e

def get_or_load_vectorstore(company_id: str, data_dir: str = "data") -> Optional[FAISS]:
    """Récupère le vectorstore depuis le cache ou le charge depuis le disque."""
    # Vérifier le cache
    if company_id in vectorstores_cache:
        return vectorstores_cache[company_id]
    
    # Essayer de charger depuis le disque
    company_index_dir = get_company_index_dir(company_id, data_dir)
    if os.path.exists(company_index_dir) and os.listdir(company_index_dir):
        try:
            embedder = OpenAIEmbeddings(model="text-embedding-3-large")
            vectordb = FAISS.load_local(company_index_dir, embedder, allow_dangerous_deserialization=True)
            vectorstores_cache[company_id] = vectordb
            return vectordb
        except Exception as e:
            print(f"Erreur lors du chargement de l'index pour l'entreprise {company_id}: {e}")
    
    return None

def get_answer(question: str,
               company_id: str,
               llm: BaseLanguageModel,
               data_dir: str = "data",
               *,
               k: int = 10,
               rerank_top_n: int = 5,
               chain_type: str = "stuff",
               temperature: float = 0.0,
               max_tokens: int = 512,
               return_sources: bool = True,
               ) -> Union[str, Dict[str, Union[str, List[Document]]]]:
    """
    Retrieve an answer to a question via a RAG pipeline with reranking for a specific company.
    """
    if not question or not question.strip():
        raise ValueError("La question est vide - veuillez fournir du texte.")

    # Récupérer le vectorstore de l'entreprise
    vectordb = get_or_load_vectorstore(company_id, data_dir)
    if vectordb is None:
        raise ValueError(f"Aucun index trouvé pour l'entreprise {company_id}. Veuillez d'abord construire l'index.")
    
    retriever = vectordb.as_retriever()

    # 1. Top-K retrieval (vector)
    retriever.search_kwargs["k"] = k

    # 2. Rerank using FlashRank (French-compatible)
    flashrank_model = "ms-marco-TinyBERT-L-2-v2" # "bce-reranker-base_v1"  # Multilingual
    client_ranker = Ranker(model_name=flashrank_model)
    compressor = FlashrankRerank(client=client_ranker, top_n=rerank_top_n)
    compression_retriever = ContextualCompressionRetriever(
        base_compressor=compressor,
        base_retriever=retriever,
    )

    # 3. Build French prompt (for both map and combine steps)
    map_template = (
        "Vous êtes un assistant client et expert en RAG. "
        "Répondez à la question du client en vous basant *uniquement* sur le contexte fourni. "
        "Si l'information n'y figure pas, dites-le clairement et n'abordez aucune discussion sans rapport avec le contexte fourni. "
        "Soyez très clair, concis et précis. Faites des réponses courtes. \n\n"
        "Contexte :\n{context}\n\n"
        "Question : {question}\n\n"
        "Réponse :"
    )
    map_prompt = PromptTemplate(template=map_template, input_variables=["context", "question"])

    # 4. Build and run the QA chain
    try:
        # Map-Reduce chains in new LangChain: prompt customization must be done via load_qa_chain (see error discussion)
        #from langchain.chains import load_qa_chain
        from langchain.chains.question_answering import load_qa_chain

        docs = compression_retriever.get_relevant_documents(question)

        qa_chain = load_qa_chain(
            llm,
            chain_type=chain_type,
            prompt=map_prompt
        )
        result = qa_chain({
            "input_documents": docs,
            "question": question
        })
    except Exception as e:
        raise RuntimeError(f"Erreur RAG : {e}")

    if return_sources:
        return {
            "answer": result["output_text"],
            "sources": docs
        }
    return result["output_text"]

def clear_company_cache(company_id: str):
    """Vide le cache vectorstore pour une entreprise spécifique."""
    if company_id in vectorstores_cache:
        del vectorstores_cache[company_id]

def get_company_stats(company_id: str, data_dir: str = "data") -> Dict:
    """Retourne des statistiques sur les documents et l'index d'une entreprise."""
    company_data_dir = get_company_data_dir(company_id, data_dir)
    company_index_dir = get_company_index_dir(company_id, data_dir)
    
    # Compter les documents
    doc_count = 0
    total_size = 0
    if os.path.exists(company_data_dir):
        for fname in os.listdir(company_data_dir):
            if fname.endswith(('.pdf', '.docx')):
                doc_count += 1
                fpath = os.path.join(company_data_dir, fname)
                total_size += os.path.getsize(fpath)
    
    # Vérifier si l'index existe
    index_exists = os.path.exists(company_index_dir) and os.listdir(company_index_dir)
    
    return {
        "company_id": company_id,
        "document_count": doc_count,
        "total_size_bytes": total_size,
        "index_exists": index_exists,
        "in_cache": company_id in vectorstores_cache
    } 